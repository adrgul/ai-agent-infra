"""
Word document parser
"""

import logging
from pathlib import Path

from langchain.schema import Document

from .base_parser import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parser for Word documents (.docx)"""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ["docx"]

    def parse(self, file_path: str) -> Document:
        """
        Parse a Word document

        Args:
            file_path: Path to the .docx file

        Returns:
            LangChain Document object
        """
        try:
            from docx import Document as DocxDocument

            path = Path(file_path)

            # Load document
            doc = DocxDocument(path)

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            # Join paragraphs
            content = "\n\n".join(paragraphs)

            # Clean the text
            cleaned_content = self._clean_text(content)

            # Extract metadata
            metadata = self.extract_metadata(file_path)
            metadata.update({
                "parser": "DocxParser",
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "word_count": len(cleaned_content.split()),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            })

            logger.info(f"Successfully parsed Word document: {file_path}")

            return Document(
                page_content=cleaned_content,
                metadata=metadata
            )

        except ImportError:
            raise RuntimeError("python-docx is required for .docx file parsing. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Failed to parse Word document {file_path}: {str(e)}")
            raise RuntimeError(f"Word document parsing failed: {str(e)}")